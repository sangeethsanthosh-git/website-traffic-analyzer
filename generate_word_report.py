from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = BASE_DIR / "WEBSITE_TRAFFIC_ANALYZER_REPORT.docx"
ARCH_IMAGE = BASE_DIR / "system_architecture.png"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_paragraph(document, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False, before=0, after=6, line=1.5):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(document, text, level_size=14):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=6, after=6, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=level_size, bold=True)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, before=0, after=3, line=1.3)
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
    else:
        run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_numbered(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, before=0, after=3, line=1.3)
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
    else:
        run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=11)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field in Microsoft Word to refresh the table of contents."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)
    set_run_font(run, size=12)


def add_page_break(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def configure_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for sec in document.sections:
        footer = sec.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_p)


def get_font(size):
    candidates = [
        "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/timesbd.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def create_architecture_image(output_path: Path):
    width, height = 1600, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(42)
    box_font = get_font(28)

    draw.text((width // 2 - 290, 40), "System Architecture - Website Traffic Analyzer", fill="black", font=title_font)

    teal = "#d9f3ef"
    orange = "#fde6d5"
    blue = "#ddeaf7"
    gray = "#eef2f7"

    boxes = {
        "user": ((670, 120, 930, 200), "User", orange),
        "browser": ((620, 250, 980, 340), "Web Browser", gray),
        "flask": ((520, 400, 1080, 510), "Flask Application (app.py)", teal),
        "auth": ((120, 610, 470, 730), "Login / Authentication\n(SQLite)", blue),
        "data": ((620, 610, 980, 730), "CSV Processing and\nAnalytics", blue),
        "chart": ((1130, 610, 1480, 730), "Chart Generation\n(Matplotlib)", blue),
        "db": ((150, 790, 440, 860), "data/app.db", gray),
        "csv": ((640, 790, 960, 860), "Traffic.csv\nTop Website Benchmark CSV", gray),
        "html": ((1120, 790, 1490, 860), "HTML Dashboard Response", gray),
    }

    for left, top, right, bottom in [value[0] for value in boxes.values()]:
        pass

    for key in boxes:
        (left, top, right, bottom), label, fill = boxes[key]
        draw.rounded_rectangle((left, top, right, bottom), radius=18, outline="black", width=3, fill=fill)
        lines = label.split("\n")
        total_height = 0
        metrics = []
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=box_font)
            metrics.append((line, bbox))
            total_height += bbox[3] - bbox[1] + 6
        current_y = top + ((bottom - top) - total_height) / 2
        for line, bbox in metrics:
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            draw.text((left + ((right - left) - text_width) / 2, current_y), line, fill="black", font=box_font)
            current_y += text_height + 6

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="black", width=3)
        draw.polygon([(x2, y2), (x2 - 12, y2 - 8), (x2 - 12, y2 + 8)], fill="black")

    arrow(800, 200, 800, 250)
    arrow(800, 340, 800, 400)
    arrow(650, 510, 300, 610)
    arrow(800, 510, 800, 610)
    arrow(950, 510, 1300, 610)
    arrow(300, 730, 300, 790)
    arrow(800, 730, 800, 790)
    arrow(1300, 730, 1300, 790)

    image.save(output_path)


def create_title_page(document: Document):
    add_paragraph(document, "WEBSITE TRAFFIC ANALYZER", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, before=80, after=18, line=1.0)
    add_paragraph(document, "A project report submitted to ICT Academy of Kerala", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=6, line=1.0)
    add_paragraph(document, "in partial fulfillment of the requirements", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=6, line=1.0)
    add_paragraph(document, "for the certification of", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=12, line=1.0)
    add_paragraph(document, "ADVANCED PYTHON", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=18, line=1.0)
    add_paragraph(document, "submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=12, line=1.0)
    add_paragraph(document, "SANGEETH", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=60, line=1.0)
    add_paragraph(document, "ICT ACADEMY OF KERALA", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, before=120, after=6, line=1.0)
    add_paragraph(document, "THIRUVANANTHAPURAM, KERALA, INDIA", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=6, line=1.0)
    add_paragraph(document, "MAY 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=0, line=1.0)


def create_list_of_figures(document: Document):
    add_heading(document, "LIST OF FIGURES", level_size=14)
    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["SL NO.", "FIGURES", "PAGE NO"]
    values = ["fig 4.1.1", "SYSTEM ARCHITECTURE", "17"]
    for index, text in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=12, bold=True)
    for index, text in enumerate(values):
        paragraph = table.rows[1].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=12)


def create_toc_page(document: Document):
    add_heading(document, "TABLE OF CONTENTS", level_size=14)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=6, after=6, line=1.3)
    add_toc(paragraph)


def add_section_title(document, text):
    paragraph = add_heading(document, text, level_size=14)
    paragraph.style = document.styles["Heading 1"]
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=True)
    return paragraph


def add_subsection_title(document, text):
    paragraph = add_heading(document, text, level_size=12)
    paragraph.style = document.styles["Heading 2"]
    for run in paragraph.runs:
        set_run_font(run, size=12, bold=True)
    return paragraph


def build_report(document: Document):
    add_section_title(document, "ABSTRACT")
    add_paragraph(document, "The Website Traffic Analyzer is a web-based analytics application developed using the Flask framework to analyze website traffic data from local datasets and present the results through summaries, charts, and benchmark comparisons. The primary objective of the system is to provide a simple and practical platform for understanding traffic behavior without depending on complex enterprise analytics tools or external online services.")
    add_paragraph(document, "The application reads traffic information from CSV files stored in the local `data` folder and transforms the raw records into meaningful insights. It calculates key metrics such as total users, average users per day, latest traffic count, highest traffic day, lowest traffic day, and overall traffic trend. In addition to daily traffic analysis, the project also includes a benchmark dataset of major websites such as google.com, youtube.com, reddit.com, and amazon.com so that users can compare global traffic performance and search for a specific website.")
    add_paragraph(document, "The project is implemented using Python, Flask, NumPy, Matplotlib, HTML, CSS, and SQLite. NumPy is used for numerical calculations, Matplotlib is used for chart generation, and SQLite is used to store user login details in SQL form. A login and registration system is included to control dashboard access and demonstrate secure data storage using hashed passwords and session-based authentication.")
    add_paragraph(document, "Overall, the Website Traffic Analyzer provides a compact and effective solution for academic demonstration and beginner-level analytics learning. It combines data analysis, visualization, SQL storage, and web development concepts in a single project that is easy to run, understand, and extend.")

    add_page_break(document)
    add_section_title(document, "1. PROBLEM DEFINITION")
    add_subsection_title(document, "1.1 OVERVIEW")
    add_paragraph(document, "Analyzing website traffic is important for understanding user engagement, measuring popularity, and identifying growth trends. In many small projects and educational settings, traffic data is available only as raw CSV files. Without a proper application, users must rely on spreadsheets or manual observation to understand the information, which is time-consuming and not visually effective.")
    add_paragraph(document, "The Website Traffic Analyzer is developed to provide a web-based interface for studying traffic data in a more meaningful and organized way. The system allows users to securely log in, view a traffic dashboard, and explore benchmark website comparisons through charts and summaries.")
    add_subsection_title(document, "1.2 PROBLEM STATEMENT")
    add_paragraph(document, "Many users face difficulty in understanding website traffic records efficiently because raw datasets alone do not provide clear insights. The major problems identified are:")
    for bullet in [
        "Difficulty in interpreting raw traffic CSV records",
        "Manual and repetitive calculations for averages and trends",
        "Lack of integrated chart-based visualization",
        "No structured benchmark comparison for well-known websites",
        "Absence of SQL-backed authentication in a basic analytics demo",
        "Limited presentation value when analysis is done only in spreadsheets",
    ]:
        add_bullet(document, bullet)
    add_paragraph(document, "To overcome these issues, the Website Traffic Analyzer is proposed as a Flask-based web application that provides secure login, traffic summaries, graphical visualization, and benchmark search in a single platform.")

    add_page_break(document)
    add_section_title(document, "2. INTRODUCTION")
    add_paragraph(document, "In the modern digital world, websites rely heavily on traffic data to understand performance and user behavior. Traffic trends help in identifying active periods, measuring reach, and planning improvements. However, raw data alone is not always useful unless it is processed and presented in a clear way.")
    add_paragraph(document, "The Website Traffic Analyzer is a web-based application developed to convert local traffic datasets into readable insights. It is designed as a small but practical project that demonstrates how Python can be used for analytics, chart generation, and web application development together. Instead of depending on online analytics platforms, the project uses local CSV files and produces its own charts and summaries.")
    add_paragraph(document, "The system also includes a login and registration mechanism backed by SQLite. This makes the project more complete by combining analytics features with SQL-based data storage. The result is a beginner-friendly application that is well suited for workshops, presentations, and academic project submission.")

    add_page_break(document)
    add_section_title(document, "3. SYSTEM ANALYSIS")
    add_subsection_title(document, "3.1 EXISTING SYSTEM")
    add_paragraph(document, "In the existing approach, traffic analysis is often performed manually using spreadsheets or static tables. These methods have several limitations:")
    for bullet in [
        "Lack of centralized dashboard-based presentation",
        "Manual effort required for totals, averages, and trends",
        "No integrated website benchmark comparison",
        "Difficulty in presenting charts within a web interface",
        "No secure user authentication in simple demo systems",
    ]:
        add_bullet(document, bullet)
    add_subsection_title(document, "3.2 PROPOSED SYSTEM")
    add_paragraph(document, "The proposed system is a Flask-based Website Traffic Analyzer that provides an organized and interactive way to study website traffic. The system includes login authentication, CSV dataset processing, summary cards, benchmark website search, and Matplotlib-based chart visualization.")
    add_paragraph(document, "Features of the proposed system are:")
    proposed = [
        "Secure user registration and login using SQLite",
        "Traffic analysis from local CSV data",
        "Summary generation for total users, average users, peak day, low day, and traffic trend",
        "Benchmark comparison for top websites worldwide",
        "Search support for benchmark websites such as reddit.com and amazon.com",
        "Graphical visualization using Matplotlib",
        "Responsive and user-friendly browser interface",
    ]
    for item in proposed:
        add_bullet(document, item)
    add_subsection_title(document, "3.3 FEASIBILITY STUDY")
    add_subsection_title(document, "3.3.1 TECHNICAL FEASIBILITY")
    add_paragraph(document, "The project is technically feasible because it uses standard and stable technologies such as Python, Flask, HTML, CSS, NumPy, Matplotlib, pytest, and SQLite. These tools are appropriate for a lightweight analytics application and can be executed on a normal system without specialized hardware or paid infrastructure.")
    add_subsection_title(document, "3.3.2 ECONOMIC FEASIBILITY")
    add_paragraph(document, "The system is economically feasible because it is built entirely using open-source technologies. No licensing cost is required for Python, Flask, SQLite, NumPy, Matplotlib, or pytest. The project can be developed and executed on a regular computer using a standard browser and Python environment.")
    add_subsection_title(document, "3.3.3 OPERATIONAL FEASIBILITY")
    add_paragraph(document, "The system is operationally feasible because it is easy to run, easy to explain, and easy to use. Users only need to start the Flask server, open the login page, sign in, and view the dashboard. This makes the project practical for labs, workshops, classroom demonstrations, and mini project evaluation.")
    add_subsection_title(document, "3.4 TECHNOLOGIES USED")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["CATEGORY", "TECHNOLOGY / TOOL", "PURPOSE"]
    rows = [
        ("Frontend", "HTML, CSS", "User interface and styling"),
        ("Backend", "Flask", "Routing, request handling, sessions, rendering"),
        ("Programming Language", "Python", "Core application logic"),
        ("Database", "SQLite", "Storage of user login details"),
        ("Python Library", "NumPy", "Numerical calculations"),
        ("Python Library", "Matplotlib", "Charts and visualization"),
        ("Testing Tool", "pytest", "Automated testing"),
        ("Version Control", "Git and GitHub", "Source control"),
    ]
    for idx, text in enumerate(headers):
        paragraph = table.rows[0].cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=11, bold=True)
    for row_data in rows:
        row = table.add_row().cells
        for idx, text in enumerate(row_data):
            paragraph = row[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_run_font(run, size=11)
    add_subsection_title(document, "3.5 LANGUAGE SPECIFICATIONS")
    add_subsection_title(document, "3.5.1 PYTHON")
    add_paragraph(document, "Python is the main programming language used in this project. It is responsible for reading CSV datasets, calculating traffic insights, generating charts, handling login logic, and managing SQLite database operations.")
    add_subsection_title(document, "3.5.2 FLASK")
    add_paragraph(document, "Flask is used as the backend framework for route handling, HTML rendering, session management, and integration of analytics results into the dashboard interface.")
    add_subsection_title(document, "3.5.3 SQLITE")
    add_paragraph(document, "SQLite is used as the local SQL database for storing user registration and login information in the file `data/app.db`. It is lightweight and well suited for a small academic project.")

    add_page_break(document)
    add_section_title(document, "4. SYSTEM DESIGN")
    add_paragraph(document, "System design defines the architecture, workflow, and main modules of the Website Traffic Analyzer. The project is designed as a web-based system where the browser interacts with the Flask backend, datasets are processed in Python, and user credentials are stored in SQLite.")
    add_paragraph(document, "The design focuses on modularity, simplicity, and clarity. The authentication flow is separated from dashboard rendering, dataset handling is organized into dedicated functions, and charts are generated dynamically before being embedded in HTML pages.")
    add_paragraph(document, "Main modules included in the system are:")
    for bullet in [
        "User Authentication Module - Handles registration, login, logout, and sessions",
        "Traffic Dataset Module - Loads and validates daily traffic records",
        "Benchmark Module - Loads website ranking data and supports search",
        "Analytics Module - Calculates totals, averages, volatility, and trend direction",
        "Visualization Module - Generates charts using Matplotlib",
        "Presentation Module - Renders templates and dashboard pages",
    ]:
        add_bullet(document, bullet)
    add_subsection_title(document, "4.1 SYSTEM ARCHITECTURE")
    add_paragraph(document, "The overall architecture of the Website Traffic Analyzer is shown below.")
    document.add_picture(str(ARCH_IMAGE), width=Inches(6.4))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("fig 4.1.1 SYSTEM ARCHITECTURE")
    set_run_font(run, size=11, bold=True)

    add_page_break(document)
    add_section_title(document, "5. PROJECT DESCRIPTION")
    add_paragraph(document, "The Website Traffic Analyzer is a web-based analytics application developed to help users understand website traffic information in a simple and structured format. The project reads local datasets, processes the data using Python, and displays the output as a dashboard with charts, benchmark tables, and quick insights.")
    add_paragraph(document, "The application has an authentication layer that restricts dashboard access to logged-in users. New users can register through the login page, and their credentials are stored in the SQLite database using hashed passwords. Once logged in, users can access the main dashboard and analyze the available datasets.")
    add_paragraph(document, "The traffic dashboard summarizes recorded days, total users, average traffic per day, and the overall trend of the dataset. In addition to the daily traffic chart, the application also provides a benchmark section that compares major websites by estimated monthly visits and allows users to search for specific domains. This improves the educational value of the project by combining local dataset analytics with broader traffic comparison.")

    add_page_break(document)
    add_section_title(document, "6. SYSTEM TESTING AND IMPLEMENTATION")
    add_subsection_title(document, "6.1 SYSTEM TESTING")
    add_paragraph(document, "System testing was carried out to verify that the application works correctly and that the modules interact properly. The following forms of testing were included:")
    testing_items = [
        "Unit testing for dataset processing, benchmark search, and summary functions",
        "Integration testing for route behavior, session flow, and SQLite-backed login",
        "Functional testing for registration, login, logout, and dashboard access",
        "User interface testing for expected page rendering and content visibility",
        "Security testing for password hashing and protected routes",
    ]
    for item in testing_items:
        add_numbered(document, item)
    add_paragraph(document, "The project currently passes the automated pytest suite successfully.")
    add_subsection_title(document, "6.2 SYSTEM IMPLEMENTATION")
    add_paragraph(document, "The system was implemented using Flask for the backend, HTML and CSS for the frontend, SQLite for user storage, NumPy for numerical analysis, and Matplotlib for chart generation. The implementation process included database setup, route creation, session handling, dataset parsing, chart rendering, and automated testing.")
    add_paragraph(document, "The application was run and verified in a browser environment. It provides a working login flow and a protected analytics dashboard suitable for demonstration and academic presentation.")

    add_page_break(document)
    add_section_title(document, "7. SYSTEM MAINTENANCE")
    add_paragraph(document, "System maintenance is required to keep the application secure, stable, and compatible with future updates. In the Website Traffic Analyzer, maintenance can include fixing bugs, improving performance, updating dependencies, and enhancing the user interface.")
    add_paragraph(document, "Types of maintenance applicable to this project are:")
    for item in [
        "Corrective maintenance for debugging route or analytics issues",
        "Adaptive maintenance for compatibility with new Python or Flask versions",
        "Perfective maintenance for interface improvements and feature upgrades",
        "Preventive maintenance for security and long-term reliability",
    ]:
        add_numbered(document, item)

    add_page_break(document)
    add_section_title(document, "8. FUTURE ENHANCEMENTS")
    add_paragraph(document, "The Website Traffic Analyzer can be extended with several advanced features in the future:")
    for item in [
        "CSV upload directly from the browser",
        "Date range filtering for traffic analysis",
        "Export of charts and reports to PDF or Excel",
        "Role-based access control for multiple user types",
        "Live traffic API integration",
        "Password reset and stronger authentication options",
        "More benchmark datasets and dynamic comparison views",
        "Cloud database support for online deployment",
    ]:
        add_bullet(document, item)

    add_page_break(document)
    add_section_title(document, "9. CONCLUSION")
    add_paragraph(document, "The Website Traffic Analyzer was successfully developed as a Flask-based web application for studying website traffic data. The system combines secure login, CSV processing, statistical analysis, chart visualization, and benchmark website comparison in a single platform.")
    add_paragraph(document, "The project demonstrates how Python can be used effectively for web development and analytics together. By integrating Flask, NumPy, Matplotlib, HTML, CSS, and SQLite, the application provides a useful and presentation-friendly solution for academic work, workshops, and learning environments.")
    add_paragraph(document, "Overall, the project achieves its objective of transforming raw traffic data into meaningful and visually understandable insights while remaining simple enough to explain and maintain.")

    add_page_break(document)
    add_section_title(document, "10. BIBILIOGRAPHY")
    bibliography = [
        "https://www.python.org/",
        "https://docs.python.org/3/",
        "https://flask.palletsprojects.com/",
        "https://numpy.org/doc/",
        "https://matplotlib.org/stable/",
        "https://www.sqlite.org/docs.html",
        "https://docs.pytest.org/",
        "https://developer.mozilla.org/",
    ]
    for idx, item in enumerate(bibliography, 1):
        add_paragraph(document, f"{idx}. {item}", align=WD_ALIGN_PARAGRAPH.LEFT, after=3, line=1.15)

    add_page_break(document)
    add_section_title(document, "11. APPENDIX")
    add_subsection_title(document, "11.1 PROJECT STRUCTURE")
    structure_lines = [
        "ictproject/",
        "    app.py",
        "    data/",
        "        Traffic.csv",
        "        top_websites_worldwide_feb_2025.csv",
        "        app.db",
        "    traffic_analyzer/",
        "        static/",
        "            style.css",
        "        templates/",
        "            base.html",
        "            index.html",
        "            login.html",
        "    tests/",
        "        test_analytics.py",
        "        test_benchmark.py",
        "        test_dataset_loader.py",
        "        test_routes.py",
        "    requirements.txt",
        "    README.md",
    ]
    for line in structure_lines:
        add_paragraph(document, line, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=0)
    add_subsection_title(document, "11.2 MAIN ROUTES")
    route_table = document.add_table(rows=1, cols=3)
    route_table.style = "Table Grid"
    route_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    route_headers = ["ROUTE", "METHOD", "PURPOSE"]
    route_rows = [
        ("/login", "GET, POST", "Login and registration page"),
        ("/logout", "POST", "Logout current user"),
        ("/", "GET", "Protected dashboard page"),
    ]
    for idx, text in enumerate(route_headers):
        paragraph = route_table.rows[0].cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=11, bold=True)
    for route_data in route_rows:
        row = route_table.add_row().cells
        for idx, text in enumerate(route_data):
            paragraph = row[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_run_font(run, size=11)


def main():
    create_architecture_image(ARCH_IMAGE)
    document = Document()
    configure_document(document)
    create_title_page(document)
    add_page_break(document)
    create_list_of_figures(document)
    add_page_break(document)
    create_toc_page(document)
    add_page_break(document)
    build_report(document)
    document.save(OUTPUT_DOCX)
    print(f"Created {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
